# ---------------------------- 更新了排班逻辑 2026.04 ----------------------------

import os
mport sys
import re
import shutil
import tkinter as tk
from tkinter import ttk, scrolledtext, messagebox, filedialog
from copy import copy
import pandas as pd
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.page import PageMargins
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

# ---------------------------- 资源路径处理（兼容 PyInstaller） ----------------------------
def resource_path(relative_path):
    """获取内置资源的绝对路径（开发环境或打包后均有效）"""
    try:
        # PyInstaller 临时解压目录
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

# ---------------------------- 目录配置（统一） ----------------------------
DESKTOP = os.path.join(os.path.expanduser("~"), "Desktop")
BASE_DIR = os.path.join(DESKTOP, "考场打印材料")
TEMPLATE_DIR = os.path.join(BASE_DIR, "模板文件")
INPUT_DIR = os.path.join(BASE_DIR, "输入文件")
OUTPUT_DIR = os.path.join(BASE_DIR, "输出结果")

# 座位标签参数（用于模板生成）
LABEL_PER_ROW = 5
LABEL_COL_SPAN = 6
LABEL_ROW_SPAN = 5
STUDENTS_PER_PAGE = 25

# 确保目录存在
def ensure_dir(directory):
    if not os.path.exists(directory):
        os.makedirs(directory)

for d in [BASE_DIR, TEMPLATE_DIR, INPUT_DIR, OUTPUT_DIR]:
    ensure_dir(d)


# ---------------------------- 排考模块（原程序1核心） ----------------------------
class ExamArrangeCore:
    """考场安排核心逻辑（从成绩文件生成学生考试信息）"""

    @staticmethod
    def generate_student_exam_info(score_file, exam_prefix, capacities, output_student_file, output_capacity_file,
                                   log_callback=None):
        """
        根据成绩文件和考场容量安排考场和座号
        返回 (成功标志, 学生列表, 错误信息)
        """
        try:
            if log_callback:
                log_callback("开始读取成绩文件...")
            df = pd.read_excel(score_file)

            # 判断排序依据（优先使用有实际数据的列）
            has_total = '总分' in df.columns
            has_rank = '年级排名' in df.columns

            # 检查总分列是否有有效数据（非空）
            if has_total:
                total_col = df['总分']
                # 去除空值后看是否还有数据
                if total_col.notna().sum() == 0:
                    has_total = False

            if not has_total and not has_rank:
                return False, None, "成绩文件中必须包含“总分”或“年级排名”列，且至少有一列有实际数据！"

            if has_total:
                rank_col = '总分'
                ascending_rank = False
                sort_msg = "总分（从高到低）"
            else:
                rank_col = '年级排名'
                ascending_rank = True
                sort_msg = "年级排名（从先到后）"

            # 检查必备列
            required = ['姓名', '班级', rank_col]
            for col in required:
                if col not in df.columns:
                    return False, None, f"成绩文件中缺少必填列：{col}"

            # 数据清洗
            df = df.dropna(subset=['姓名', '班级'])
            total_students = len(df)
            if total_students == 0:
                return False, None, "成绩文件中没有有效的学生数据"

            # 直接按排名列排序（纯年级排名顺序）
            df_sorted = df.sort_values([rank_col, '姓名'], ascending=[ascending_rank, True])

            fixed_capacity = sum(capacities)
            last_court_capacity = total_students - fixed_capacity
            if last_court_capacity < 0:
                return False, None, f"前{len(capacities)}个考场总容量({fixed_capacity})已超过学生总数({total_students})！"

            all_caps = capacities.copy()
            has_auto_court = False
            if last_court_capacity > 0:
                all_caps.append(last_court_capacity)
                has_auto_court = True

            # 分配考场和座号
            allocations = []
            court_idx = 0
            seat_in_court = 0
            for idx, (_, student) in enumerate(df_sorted.iterrows()):
                if seat_in_court >= all_caps[court_idx]:
                    court_idx += 1
                    seat_in_court = 0
                exam_id = f"{exam_prefix}{(idx + 1):03d}"
                allocations.append({
                    '准考证号': exam_id,
                    '姓名': student['姓名'],
                    '班级': int(student['班级']),
                    '考场号': court_idx + 1,
                    '座号': seat_in_court + 1
                })
                seat_in_court += 1

            result_df = pd.DataFrame(allocations)
            result_df = result_df.sort_values(['考场号', '座号'])
            result_df.to_excel(output_student_file, index=False)

            # 生成考场容量汇总表
            court_info = []
            manual_court_count = len(capacities)
            for court_num in range(1, len(all_caps) + 1):
                court_students = result_df[result_df['考场号'] == court_num]
                student_count = len(court_students)
                start_id = court_students.iloc[0]['准考证号'] if student_count > 0 else ''
                end_id = court_students.iloc[-1]['准考证号'] if student_count > 0 else ''
                set_cap = all_caps[court_num - 1] if court_num <= manual_court_count else '自动计算'
                court_info.append({
                    '考场号': court_num,
                    '设置容量': set_cap,
                    '实际人数': student_count,
                    '起始考号': start_id,
                    '截止考号': end_id
                })
            capacity_df = pd.DataFrame(court_info)
            capacity_df.to_excel(output_capacity_file, index=False)

            # 返回学生列表（用于后续打印材料生成）
            student_list = result_df.to_dict('records')
            # 格式化班级、考场号、座号为字符串（便于后续使用）
            for stu in student_list:
                stu['班级'] = str(stu['班级']).zfill(2)
                stu['考场号'] = str(stu['考场号']).zfill(2)
                stu['座号'] = str(stu['座号']).zfill(2)

            if log_callback:
                log_callback(f"排序依据：{sort_msg}，学生总数：{total_students}人，考场总数：{len(all_caps)}个")
                log_callback(f"学生考试信息已保存：{output_student_file}")
                log_callback(f"考场容量汇总已保存：{output_capacity_file}")

            return True, student_list, None
        except Exception as e:
            return False, None, str(e)


# ---------------------------- 打印材料生成模块（原程序2核心） ----------------------------
class PrintMaterialCore:
    """生成门贴、座位标签、班级考场信息"""

    @staticmethod
    def find_template_file(template_name):
        """在模板目录中查找文件，若不存在则从内置资源复制，再失败则创建默认模板"""
        template_path = os.path.join(TEMPLATE_DIR, template_name)
        if not os.path.exists(template_path):
            # 从内置资源复制
            src = resource_path(os.path.join("内置模板", template_name))
            if os.path.exists(src):
                shutil.copy2(src, template_path)
            else:
                # 内置资源也不存在，则动态生成
                if template_name == "门贴模板.docx":
                    PrintMaterialCore._create_default_door_template()
                elif template_name == "座位标签模板.xlsx":
                    PrintMaterialCore._create_default_label_template()
        return template_path if os.path.exists(template_path) else None

    @staticmethod
    def _create_default_door_template():
        """生成默认的门贴模板（备用）"""
        path = os.path.join(TEMPLATE_DIR, "门贴模板.docx")
        doc = Document()
        for section in doc.sections:
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run("«考试标题»")
        title_run.font.size = Pt(24)
        title_run.font.bold = True
        doc.add_paragraph()
        table = doc.add_table(rows=4, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.cell(0, 0).text = "考试地点："
        table.cell(0, 1).text = "«考试地点»"
        table.cell(0, 2).text = "«考试地点»"
        table.cell(1, 0).text = "考场号："
        table.cell(1, 1).text = "«考场号»"
        table.cell(1, 2).text = "«考场号»"
        table.cell(3, 0).text = "«起始编号1»"
        table.cell(3, 1).text = "---"
        table.cell(3, 2).text = "«结束编号1»"
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.size = Pt(18)
                        run.font.bold = True
        doc.save(path)

    @staticmethod
    def _create_default_label_template():
        """生成默认座位标签模板（5×5网格，备用）"""
        path = os.path.join(TEMPLATE_DIR, "座位标签模板.xlsx")
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "座位标签"
        ws.page_setup.orientation = ws.ORIENTATION_PORTRAIT
        ws.page_setup.paperSize = ws.PAPERSIZE_A4
        ws.page_margins = PageMargins(left=0.2, right=0.2, top=0.2, bottom=0.2)

        title_font = Font(bold=True, size=12, name="宋体")
        label_font = Font(size=11, name="宋体")
        center_align = Alignment(horizontal='center', vertical='center')

        for row_group in range(5):
            base_row = row_group * LABEL_ROW_SPAN + 1
            for col_group in range(LABEL_PER_ROW):
                base_col = col_group * LABEL_COL_SPAN + 1
                ws.cell(row=base_row, column=base_col, value="第二实验中学").font = title_font
                ws.cell(row=base_row, column=base_col).alignment = center_align
                ws.cell(row=base_row + 1, column=base_col, value="准考证号：").font = label_font
                ws.cell(row=base_row + 1, column=base_col).alignment = center_align
                ws.cell(row=base_row + 2, column=base_col, value="姓名：").font = label_font
                ws.cell(row=base_row + 2, column=base_col).alignment = center_align
                ws.cell(row=base_row + 2, column=base_col + 3, value="班级：").font = label_font
                ws.cell(row=base_row + 2, column=base_col + 3).alignment = center_align
                ws.cell(row=base_row + 3, column=base_col, value="考场：").font = label_font
                ws.cell(row=base_row + 3, column=base_col).alignment = center_align
                ws.cell(row=base_row + 3, column=base_col + 3, value="座号：").font = label_font
                ws.cell(row=base_row + 3, column=base_col + 3).alignment = center_align
                ws.cell(row=base_row + 4, column=base_col).alignment = center_align

        for col in range(1, LABEL_PER_ROW * LABEL_COL_SPAN + 1):
            ws.column_dimensions[get_column_letter(col)].width = 3.8
        for row in range(1, 5 * LABEL_ROW_SPAN + 1):
            ws.row_dimensions[row].height = 22
        wb.save(path)

    @staticmethod
    def _apply_font_style(target_run, source_run):
        if source_run is None:
            return
        target_run.font.name = source_run.font.name
        target_run.font.size = source_run.font.size
        target_run.font.bold = source_run.font.bold
        target_run.font.italic = source_run.font.italic
        target_run.font.underline = source_run.font.underline
        if source_run.font.color and source_run.font.color.rgb:
            target_run.font.color.rgb = source_run.font.color.rgb

    @staticmethod
    def _replace_doc_placeholders(doc, replace_dict):
        pattern = re.compile(r'«(.*?)»')
        for paragraph in doc.paragraphs:
            full_text = paragraph.text
            if not pattern.search(full_text):
                continue
            original_format = None
            original_alignment = paragraph.alignment
            if paragraph.runs:
                original_format = paragraph.runs[0]
            for key, value in replace_dict.items():
                full_text = full_text.replace(key, value)
            paragraph.clear()
            new_run = paragraph.add_run(full_text)
            if original_format:
                PrintMaterialCore._apply_font_style(new_run, original_format)
            paragraph.alignment = original_alignment
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        full_text = paragraph.text
                        if not pattern.search(full_text):
                            continue
                        original_format = None
                        original_alignment = paragraph.alignment
                        if paragraph.runs:
                            original_format = paragraph.runs[0]
                        for key, value in replace_dict.items():
                            full_text = full_text.replace(key, value)
                        paragraph.clear()
                        new_run = paragraph.add_run(full_text)
                        if original_format:
                            PrintMaterialCore._apply_font_style(new_run, original_format)
                        paragraph.alignment = original_alignment

    @staticmethod
    def _set_all_text_bold_heiti(doc):
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.font.name = "微软雅黑"
                run.font.bold = True
                if run._element.rPr is not None:
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = "微软雅黑"
                            run.font.bold = True
                            if run._element.rPr is not None:
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")

    @staticmethod
    def _locate_label_positions(ws):
        """定位模板中每个标签的数据单元格位置"""
        positions = []
        max_row = ws.max_row
        max_col = ws.max_column
        for row in range(1, max_row + 1):
            for col in range(1, max_col + 1):
                cell = ws.cell(row, col)
                if cell.value and "准考证号：" in str(cell.value):
                    examid_cell = ws.cell(row, col + 2)
                    name_cell = None
                    class_cell = None
                    for r_offset in range(1, 4):
                        test_row = row + r_offset
                        for c_offset in range(0, 5):
                            test_cell = ws.cell(test_row, col + c_offset)
                            if test_cell.value and "姓名：" in str(test_cell.value):
                                name_cell = ws.cell(test_row, test_cell.column + 1)
                            if test_cell.value and "班级：" in str(test_cell.value):
                                class_cell = ws.cell(test_row, test_cell.column + 1)
                    room_cell = None
                    seat_cell = None
                    for r_offset in range(2, 5):
                        test_row = row + r_offset
                        for c_offset in range(0, 5):
                            test_cell = ws.cell(test_row, col + c_offset)
                            if test_cell.value and "考场：" in str(test_cell.value):
                                room_cell = ws.cell(test_row, test_cell.column + 1)
                            if test_cell.value and "座号：" in str(test_cell.value):
                                seat_cell = ws.cell(test_row, test_cell.column + 1)
                    if examid_cell and name_cell and class_cell and room_cell and seat_cell:
                        positions.append({
                            'examid': examid_cell,
                            'name': name_cell,
                            'class': class_cell,
                            'room': room_cell,
                            'seat': seat_cell
                        })
                    else:
                        name_cell = ws.cell(row + 2, col + 1)
                        class_cell = ws.cell(row + 2, col + 4)
                        room_cell = ws.cell(row + 3, col + 1)
                        seat_cell = ws.cell(row + 3, col + 4)
                        positions.append({
                            'examid': examid_cell,
                            'name': name_cell,
                            'class': class_cell,
                            'room': room_cell,
                            'seat': seat_cell
                        })
        return positions

    @staticmethod
    def _copy_page_style(src_ws, dst_ws, start_row, rows_per_page):
        for col in range(1, LABEL_PER_ROW * LABEL_COL_SPAN + 1):
            dst_ws.column_dimensions[get_column_letter(col)].width = src_ws.column_dimensions[
                get_column_letter(col)].width
        for r in range(1, rows_per_page + 1):
            dst_ws.row_dimensions[start_row + r - 1].height = src_ws.row_dimensions[r].height
            for c in range(1, LABEL_PER_ROW * LABEL_COL_SPAN + 1):
                src_cell = src_ws.cell(row=r, column=c)
                dst_cell = dst_ws.cell(row=start_row + r - 1, column=c)
                if src_cell.value:
                    dst_cell.value = src_cell.value
                if src_cell.has_style:
                    dst_cell.font = copy(src_cell.font)
                    dst_cell.alignment = copy(src_cell.alignment)
                    dst_cell.border = copy(src_cell.border)
                    dst_cell.fill = copy(src_cell.fill)
        for merged in list(src_ws.merged_cells.ranges):
            if merged.min_row <= rows_per_page:
                try:
                    dst_ws.merge_cells(start_row=start_row + merged.min_row - 1, start_column=merged.min_col,
                                       end_row=start_row + merged.max_row - 1, end_column=merged.max_col)
                except:
                    pass

    @staticmethod
    def _fill_label_template(ws, students):
        if not students:
            return
        positions = PrintMaterialCore._locate_label_positions(ws)
        if not positions:
            raise ValueError("无法从模板中定位标签填充位置")
        labels_per_page = len(positions)
        total_students = len(students)
        pages_needed = (total_students + labels_per_page - 1) // labels_per_page
        rows_per_page = LABEL_ROW_SPAN * 5

        for page in range(1, pages_needed):
            start_row = page * rows_per_page + 1
            PrintMaterialCore._copy_page_style(ws, ws, start_row, rows_per_page)

        for page in range(pages_needed):
            start_idx = page * labels_per_page
            end_idx = min(start_idx + labels_per_page, total_students)
            page_students = students[start_idx:end_idx]
            base_row_offset = page * rows_per_page
            for i, student in enumerate(page_students):
                if i >= labels_per_page:
                    break
                pos = positions[i]
                new_examid = ws.cell(row=pos['examid'].row + base_row_offset, column=pos['examid'].column)
                new_name = ws.cell(row=pos['name'].row + base_row_offset, column=pos['name'].column)
                new_class = ws.cell(row=pos['class'].row + base_row_offset, column=pos['class'].column)
                new_room = ws.cell(row=pos['room'].row + base_row_offset, column=pos['room'].column)
                new_seat = ws.cell(row=pos['seat'].row + base_row_offset, column=pos['seat'].column)
                new_examid.value = student['准考证号']
                new_name.value = student['姓名']
                new_class.value = student['班级']
                new_room.value = student['考场号']
                new_seat.value = student['座号']

    @staticmethod
    def read_door_data(file_path):
        wb = openpyxl.load_workbook(file_path, data_only=True)
        ws = wb.active
        doors = []
        for row in ws.iter_rows(min_row=2, values_only=True):
            if row[0] and row[1]:
                exam_id = str(row[1]).strip().zfill(2)
                doors.append((str(row[0]).strip(), exam_id))
        return doors

    @staticmethod
    def read_student_range(file_path):
        wb = openpyxl.load_workbook(file_path, data_only=True)
        ws = wb.active
        range_dict = {}
        for row in ws.iter_rows(min_row=2, values_only=True):
            if row[0] and row[3] and row[4]:
                exam_id = str(row[3]).strip().zfill(2)
                stu_id = str(row[0]).strip()
                if exam_id not in range_dict:
                    range_dict[exam_id] = []
                range_dict[exam_id].append(stu_id)
        result = {}
        for exam_id, ids in range_dict.items():
            ids_sorted = sorted(ids)
            result[exam_id] = (ids_sorted[0], ids_sorted[-1])
        return result

    @staticmethod
    def generate_door_placards(door_list, student_range, output_dir, title_text, log_callback=None):
        template_path = PrintMaterialCore.find_template_file("门贴模板.docx")
        if not template_path:
            raise FileNotFoundError("未找到门贴模板文件 '门贴模板.docx'，请放入模板文件夹。")
        total = len(door_list)
        for idx, (location, exam_id) in enumerate(door_list):
            doc = Document(template_path)
            replace_dict = {
                '«考试标题»': title_text,
                '«考试地点»': location,
                '«考场号»': exam_id,
                '«起始编号1»': '',
                '«结束编号1»': ''
            }
            if exam_id in student_range:
                replace_dict['«起始编号1»'] = student_range[exam_id][0]
                replace_dict['«结束编号1»'] = student_range[exam_id][1]
            PrintMaterialCore._replace_doc_placeholders(doc, replace_dict)
            PrintMaterialCore._set_all_text_bold_heiti(doc)
            filename = f"门贴_考场{exam_id}.docx"
            filepath = os.path.join(output_dir, filename)
            doc.save(filepath)
            if log_callback:
                log_callback(f"已生成门贴：{filepath}")

    @staticmethod
    def generate_seat_labels(student_list, output_dir, log_callback=None):
        """按考场生成座位标签"""
        template_path = PrintMaterialCore.find_template_file("座位标签模板.xlsx")
        if not template_path:
            raise FileNotFoundError("未找到座位标签模板文件 '座位标签模板.xlsx'，请放入模板文件夹。")
        room_groups = {}
        for stu in student_list:
            room = stu['考场号']
            if room not in room_groups:
                room_groups[room] = []
            room_groups[room].append(stu)
        total_rooms = len(room_groups)
        for idx, (room_name, students) in enumerate(room_groups.items()):
            if log_callback:
                log_callback(f"正在处理 考场{room_name}（{len(students)}人）...")
            students_sorted = sorted(students, key=lambda x: int(x['座号']))
            wb = openpyxl.load_workbook(template_path)
            ws = wb.active
            PrintMaterialCore._fill_label_template(ws, students_sorted)
            filename = f"座位标签_考场{str(room_name).zfill(2)}.xlsx"
            filepath = os.path.join(output_dir, filename)
            wb.save(filepath)
            if log_callback:
                log_callback(f"已生成座位标签：{filepath}")

    @staticmethod
    def generate_class_info(student_list, output_dir, title_text, log_callback=None):
        """生成班级考场信息Excel"""
        from openpyxl import Workbook
        class_groups = {}
        for stu in student_list:
            cls = stu['班级']
            if cls not in class_groups:
                class_groups[cls] = []
            class_groups[cls].append(stu)
        output_path = os.path.join(output_dir, "班级考场信息.xlsx")
        wb = Workbook()
        wb.remove(wb.active)
        total_classes = len(class_groups)
        for idx, (class_name, students) in enumerate(class_groups.items()):
            students_sorted = sorted(students, key=lambda x: (x['考场号'], int(x['座号'])))
            data = []
            for stu in students_sorted:
                data.append([stu['准考证号'], stu['姓名'], stu['班级'], stu['考场号'], stu['座号']])
            sheet_name = f"{class_name}"
            if len(sheet_name) > 31:
                sheet_name = sheet_name[:31]
            ws = wb.create_sheet(sheet_name)
            title = f"{title_text} {class_name} 班考生信息"
            ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=5)
            title_cell = ws.cell(row=1, column=1, value=title)
            title_cell.font = Font(bold=True, size=14)
            title_cell.alignment = Alignment(horizontal='center', vertical='center')
            headers = ['准考证号', '姓名', '班级', '考场号', '座号']
            for col, header in enumerate(headers, 1):
                cell = ws.cell(row=2, column=col, value=header)
                cell.font = Font(bold=True)
                cell.alignment = Alignment(horizontal='center', vertical='center')
            for row_idx, row_data in enumerate(data, start=3):
                for col_idx, value in enumerate(row_data, 1):
                    cell = ws.cell(row=row_idx, column=col_idx, value=value)
                    cell.alignment = Alignment(horizontal='center', vertical='center')
            ws.column_dimensions['A'].width = 20
            ws.column_dimensions['B'].width = 15
            ws.column_dimensions['C'].width = 10
            ws.column_dimensions['D'].width = 10
            ws.column_dimensions['E'].width = 10
            ws.row_dimensions[1].height = 25
            ws.row_dimensions[2].height = 15
            for row in range(3, len(data) + 3):
                ws.row_dimensions[row].height = 15
            thin_border = Border(left=Side(style='thin'), right=Side(style='thin'),
                                 top=Side(style='thin'), bottom=Side(style='thin'))
            last_row = len(data) + 2
            for row in range(2, last_row + 1):
                for col in range(1, 6):
                    ws.cell(row=row, column=col).border = thin_border
            ws.page_margins = PageMargins(left=0.197, right=0.197, top=0.3937, bottom=0.3937, header=0.197,
                                          footer=0.197)
            ws.page_setup.fitToWidth = 1
            ws.page_setup.fitToHeight = 1
            ws.page_setup.fitToPage = True
            ws.print_options.horizontalCentered = True
            ws.print_options.verticalCentered = True
        wb.save(output_path)
        if log_callback:
            log_callback(f"已生成班级考场信息：{output_path}")


# ---------------------------- 示例文件生成（动态生成，保持原样） ----------------------------
def create_example_door_file():
    path = os.path.join(INPUT_DIR, "考场地点数据.xlsx")
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "考场地点"
    headers = ["考场地点", "考场号"]
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = Font(bold=True, size=12)
        cell.alignment = Alignment(horizontal='center', vertical='center')
    examples = [["教学楼A栋101", "01"], ["教学楼A栋102", "02"]]
    for row, example in enumerate(examples, 2):
        for col, value in enumerate(example, 1):
            ws.cell(row=row, column=col, value=value)
    wb.save(path)

def create_example_score_template():
    """生成成绩模板（用于排考）"""
    path = os.path.join(INPUT_DIR, "学生原始成绩.xlsx")
    df_score = pd.DataFrame({
        "姓名": ["张三", "李四", "王五", "赵六"],
        "班级": [1, 1, 2, 2],
        "总分": [650, 620, 600, 580],
        "年级排名": [10, 25, 30, 45]
    })
    with pd.ExcelWriter(path, engine='openpyxl') as writer:
        df_score.to_excel(writer, index=False, sheet_name="成绩")
        worksheet = writer.sheets["成绩"]
        worksheet.cell(row=1, column=5, value="提示：总分和年级排名至少填写一项，系统将自动识别。")
        worksheet.cell(row=2, column=5, value="（请勿删除或修改列名，班级必须为数字）")

def create_example_capacity_template():
    """生成考场容量模板"""
    path = os.path.join(INPUT_DIR, "考场容量模板.xlsx")
    df_cap = pd.DataFrame({
        "考场": [1, 2],
        "容量": [2, 2]
    })
    with pd.ExcelWriter(path, engine='openpyxl') as writer:
        df_cap.to_excel(writer, index=False, sheet_name="容量")
        worksheet = writer.sheets["容量"]
        worksheet.cell(row=1, column=3, value="⚠️  核心规则说明")
        worksheet.cell(row=2, column=3, value="1. 请在此填写前N个考场的固定容量")
        worksheet.cell(row=3, column=3, value="2. 最终的最后一个考场，无需在此填写，程序将根据学生总数自动计算")
        worksheet.cell(row=4, column=3, value="3. 请勿修改「考场」「容量」列名，仅修改容量数值即可")
        last_data_row = 3
        tip_cell = worksheet.cell(row=last_data_row, column=3,
                                  value="★ 模板手动考场到此结束，无需往下新增行，最后一个考场自动计算")
        tip_cell.font = Font(color="FF0000", bold=True)
        tip_cell.fill = PatternFill(start_color="FFFF99", end_color="FFFF99", fill_type="solid")


# ---------------------------- 统一GUI界面 ----------------------------
class UnifiedExamTool:
    def __init__(self, root):
        self.root = root
        self.root.title("沂水县第二实验中学考场编排全流程工具")
        self.root.geometry("950x850")

        # 定义固定文件路径
        self.default_score_file = os.path.join(INPUT_DIR, "学生原始成绩.xlsx")
        self.default_capacity_template = os.path.join(INPUT_DIR, "考场容量模板.xlsx")
        self.default_door_file = os.path.join(INPUT_DIR, "考场地点数据.xlsx")
        self.student_info_output = os.path.join(INPUT_DIR, "学生考试信息.xlsx")  # 排考输出，也是材料生成的输入
        self.capacity_output = os.path.join(OUTPUT_DIR, "考场容量.xlsx")

        # 变量
        self.score_file = tk.StringVar(value=self.default_score_file)
        self.capacity_template = tk.StringVar(value=self.default_capacity_template)
        self.door_file = tk.StringVar(value=self.default_door_file)
        self.output_dir = tk.StringVar(value=OUTPUT_DIR)
        self.exam_prefix = tk.StringVar(value="701020")
        self.exam_title = tk.StringVar(value="沂水县第二实验中学单元作业")

        # 预览变量
        self.total_students_var = tk.StringVar(value="未统计")
        self.last_court_capacity_var = tk.StringVar(value="未计算")

        # 考场容量动态行存储
        self.capacity_rows = []  # (frame, entry, delete_btn, court_num)

        self.create_widgets()

        # 初始化模板和示例文件（如果不存在）
        self.init_default_files()

        # 默认添加3个考场（示例）
        for i in range(1, 3):
            self.add_capacity_row(i)

        # 启动时显示使用说明对话框（6秒自动关闭）
        self.show_startup_dialog()

    def show_startup_dialog(self):
        """显示程序使用说明对话框，6秒后自动关闭"""
        dialog = tk.Toplevel(self.root)
        dialog.title("使用说明与注意事项")
        dialog.geometry("500x400")
        dialog.transient(self.root)  # 设为父窗口的临时窗口
        dialog.grab_set()  # 模态

        # 文本内容
        text = tk.Text(dialog, wrap=tk.WORD, font=("微软雅黑", 10), padx=10, pady=10)
        text.pack(fill=tk.BOTH, expand=True)

        content = """
【沂水县第二实验中学考场编排工具 使用说明】

一、准备工作
1. 程序启动后会自动在桌面生成“考场打印材料”文件夹，内含：
   - 模板文件/      （存放门贴模板、座位标签模板）
   - 输入文件/      （存放学生原始成绩、考场容量模板、考场地点数据）
   - 输出结果/      （生成的门贴、座位标签、班级考场信息等）

2. 首次使用建议点击【生成示例文件】按钮，程序会在“输入文件”文件夹中生成示例：
   - 学生原始成绩.xlsx
   - 考场容量模板.xlsx
   - 考场地点数据.xlsx

二、操作流程
1. 排考设置（选项卡1）
   - 选择成绩文件（必须包含：姓名、班级、总分或年级排名）
   - 设置考号前缀（如701020，最终考号格式：701020001）
   - 设置考场容量：添加多个考场，填写每个考场的容量（最后一个考场会自动计算）
   - 可导入/导出容量模板
   - 点击【刷新统计】预览学生总数和最后一个考场容量

2. 打印材料设置（选项卡2）
   - 选择门贴数据源（考场地点数据.xlsx，格式：考场地点、考场号）
   - 填写考试标题（用于门贴和班级信息表）
   - 设置输出目录（默认为“输出结果”文件夹）

3. 执行与日志（选项卡3）
   - 【一键生成全部材料】推荐：先排考，再生成门贴、座位标签、班级信息表
   - 【仅生成排考】：只生成学生考试信息.xlsx
   - 【仅生成打印材料】：基于已有的学生考试信息.xlsx生成打印材料
   - 【清空日志】清除日志窗口内容

三、重要注意事项
1. 成绩文件必须包含“姓名”、“班级”列，以及“总分”或“年级排名”之一。
2. 班级必须是数字（如1,2,3），程序会自动补零（01,02...）。
3. 考场容量模板：仅填写前N个固定考场的容量，最后一个考场会根据学生总数自动计算，无需填写。
4. 门贴数据源中的考场号必须与学生考试信息中的考场号一致（建议用两位数字，如01,02）。
5. 座位标签模板和门贴模板支持自定义，放在“模板文件”文件夹中，程序会自动读取。
6. 所有生成的文件均为独立文件，不会覆盖原始数据，请放心使用。

四、常见问题
- 如果生成失败，请检查日志中的错误提示。
- 若提示缺少列，请确认Excel列名是否完全匹配（总分、年级排名等）。
- 座位标签排版异常时，可替换“模板文件/座位标签模板.xlsx”自定义布局。

提示：本对话框会在30秒后自动关闭。
        """
        text.insert(tk.END, content)
        text.config(state=tk.DISABLED)

        # 按钮框架
        btn_frame = ttk.Frame(dialog)
        btn_frame.pack(pady=10)

        def close_dialog():
            dialog.destroy()

        ok_btn = ttk.Button(btn_frame, text="确定", command=close_dialog)
        ok_btn.pack()

        # 6秒后自动关闭
        self.root.after(30000, close_dialog)

    def init_default_files(self):
        """生成必要的模板和示例文件"""
        if not os.path.exists(self.default_score_file):
            create_example_score_template()
        if not os.path.exists(self.default_capacity_template):
            create_example_capacity_template()
        if not os.path.exists(self.default_door_file):
            create_example_door_file()
        # 确保座位标签模板存在（会自动从内置资源复制）
        PrintMaterialCore.find_template_file("座位标签模板.xlsx")
        PrintMaterialCore.find_template_file("门贴模板.docx")

    def create_widgets(self):
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.pack(fill=tk.BOTH, expand=True)

        # 创建笔记本（选项卡）
        notebook = ttk.Notebook(main_frame)
        notebook.pack(fill=tk.BOTH, expand=True)

        # ==================== 选项卡1：排考设置 ====================
        arrange_frame = ttk.Frame(notebook, padding="15")
        notebook.add(arrange_frame, text="1. 排考设置（生成学生考试信息）")

        # 成绩文件
        row0 = ttk.Frame(arrange_frame)
        row0.pack(fill=tk.X, pady=5)
        ttk.Label(row0, text="成绩文件（需包含姓名、班级、总分或年级排名）:", width=32, anchor=tk.W).pack(side=tk.LEFT)
        f0 = ttk.Frame(row0)
        f0.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        ttk.Entry(f0, textvariable=self.score_file).pack(side=tk.LEFT, fill=tk.X, expand=True)
        ttk.Button(f0, text="浏览",
                   command=lambda: self.browse_file(self.score_file, [("Excel文件", "*.xlsx *.xls")])).pack(
            side=tk.LEFT, padx=5)

        # 考号前缀
        row1 = ttk.Frame(arrange_frame)
        row1.pack(fill=tk.X, pady=5)
        ttk.Label(row1, text="考号前缀:", width=32, anchor=tk.W).pack(side=tk.LEFT)
        f1 = ttk.Frame(row1)
        f1.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        ttk.Entry(f1, textvariable=self.exam_prefix, width=20).pack(side=tk.LEFT)
        ttk.Label(f1, text=" (例:701020 → 701020001)").pack(side=tk.LEFT, padx=5)

        # 考场容量模板路径
        row_cap_path = ttk.Frame(arrange_frame)
        row_cap_path.pack(fill=tk.X, pady=5)
        ttk.Label(row_cap_path, text="考场容量模板路径:", width=32, anchor=tk.W).pack(side=tk.LEFT)
        f_cap_path = ttk.Frame(row_cap_path)
        f_cap_path.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        ttk.Entry(f_cap_path, textvariable=self.capacity_template).pack(side=tk.LEFT, fill=tk.X, expand=True)
        ttk.Button(f_cap_path, text="浏览",
                   command=lambda: self.browse_file(self.capacity_template, [("Excel文件", "*.xlsx *.xls")])).pack(
            side=tk.LEFT, padx=5)

        # 考场容量动态设置区域
        cap_label = ttk.Label(arrange_frame, text="考场容量设置（仅填写固定考场，最后一个自动计算）:", anchor=tk.W)
        cap_label.pack(fill=tk.X, pady=(10, 5))

        # 容量操作按钮
        cap_btn_frame = ttk.Frame(arrange_frame)
        cap_btn_frame.pack(fill=tk.X, pady=5)
        ttk.Button(cap_btn_frame, text="添加考场", command=self.add_capacity_row).pack(side=tk.LEFT, padx=2)
        ttk.Button(cap_btn_frame, text="从模板导入", command=self.import_capacity_from_template).pack(side=tk.LEFT,
                                                                                                      padx=2)
        ttk.Button(cap_btn_frame, text="导出当前为模板", command=self.export_capacity_template).pack(side=tk.LEFT,
                                                                                                     padx=2)

        # 容量滚动区域
        cap_canvas = tk.Canvas(arrange_frame, borderwidth=0, highlightthickness=0, height=200)
        cap_scrollbar = ttk.Scrollbar(arrange_frame, orient="vertical", command=cap_canvas.yview)
        self.capacity_container = ttk.Frame(cap_canvas)
        self.capacity_container.bind("<Configure>", lambda e: cap_canvas.configure(scrollregion=cap_canvas.bbox("all")))
        cap_canvas.create_window((0, 0), window=self.capacity_container, anchor="nw")
        cap_canvas.configure(yscrollcommand=cap_scrollbar.set)
        cap_canvas.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, pady=5)
        cap_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        # 预览区域（改为上下布局，避免溢出）
        preview_frame = ttk.LabelFrame(arrange_frame, text="预览信息（点击刷新按钮更新）", padding="8")
        preview_frame.pack(fill=tk.X, pady=10)

        # 第一行：学生总计人数
        preview_row1 = ttk.Frame(preview_frame)
        preview_row1.pack(fill=tk.X, pady=3)
        ttk.Label(preview_row1, text="学生总计人数:", width=20, anchor=tk.W).pack(side=tk.LEFT)
        ttk.Entry(preview_row1, textvariable=self.total_students_var, state="readonly", width=15).pack(side=tk.LEFT,
                                                                                                       padx=5)

        # 第二行：最后一个考场容量
        preview_row2 = ttk.Frame(preview_frame)
        preview_row2.pack(fill=tk.X, pady=3)
        ttk.Label(preview_row2, text="最后一个考场容量（自动计算）:", width=20, anchor=tk.W).pack(side=tk.LEFT)
        ttk.Entry(preview_row2, textvariable=self.last_court_capacity_var, state="readonly", width=15).pack(
            side=tk.LEFT, padx=5)

        # 第三行：刷新按钮和提示
        preview_row3 = ttk.Frame(preview_frame)
        preview_row3.pack(fill=tk.X, pady=3)
        ttk.Button(preview_row3, text="刷新统计", command=self.refresh_preview).pack(side=tk.LEFT, padx=5)
        ttk.Label(preview_row3, text="注：需先保存考场容量设置，点击刷新即可预览", foreground="gray").pack(side=tk.LEFT,
                                                                                                         padx=10)

        # ==================== 选项卡2：打印材料设置 ====================
        print_frame = ttk.Frame(notebook, padding="15")
        notebook.add(print_frame, text="2. 打印材料设置（门贴/座位标签/班级信息）")

        # 门贴数据源
        row_door = ttk.Frame(print_frame)
        row_door.pack(fill=tk.X, pady=5)
        ttk.Label(row_door, text="门贴数据源（考场地点与考场号对应）:", width=32, anchor=tk.W).pack(side=tk.LEFT)
        f_door = ttk.Frame(row_door)
        f_door.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        ttk.Entry(f_door, textvariable=self.door_file).pack(side=tk.LEFT, fill=tk.X, expand=True)
        ttk.Button(f_door, text="浏览",
                   command=lambda: self.browse_file(self.door_file, [("Excel文件", "*.xlsx *.xls")])).pack(side=tk.LEFT,
                                                                                                           padx=5)

        # 考试标题
        row_title = ttk.Frame(print_frame)
        row_title.pack(fill=tk.X, pady=5)
        ttk.Label(row_title, text="考试标题（用于门贴和班级信息表）:", width=32, anchor=tk.W).pack(side=tk.LEFT)
        f_title = ttk.Frame(row_title)
        f_title.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        ttk.Entry(f_title, textvariable=self.exam_title, width=50).pack(side=tk.LEFT)

        # 输出目录
        row_out = ttk.Frame(print_frame)
        row_out.pack(fill=tk.X, pady=5)
        ttk.Label(row_out, text="输出目录:", width=32, anchor=tk.W).pack(side=tk.LEFT)
        f_out = ttk.Frame(row_out)
        f_out.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        ttk.Entry(f_out, textvariable=self.output_dir).pack(side=tk.LEFT, fill=tk.X, expand=True)
        ttk.Button(f_out, text="浏览", command=lambda: self.browse_dir(self.output_dir)).pack(side=tk.LEFT, padx=5)

        # ==================== 选项卡3：执行与日志 ====================
        log_frame = ttk.Frame(notebook, padding="15")
        notebook.add(log_frame, text="3. 执行与日志")

        # 按钮区域（一键生成在最左侧，清空日志在最右侧）
        action_frame = ttk.Frame(log_frame)
        action_frame.pack(fill=tk.X, pady=10)

        # 左侧按钮组
        left_btn_frame = ttk.Frame(action_frame)
        left_btn_frame.pack(side=tk.LEFT)
        ttk.Button(left_btn_frame, text="一键生成全部材料", command=self.run_full, style="Accent.TButton").pack(
            side=tk.LEFT, padx=5)
        ttk.Button(left_btn_frame, text="仅生成排考", command=self.run_arrange_only).pack(side=tk.LEFT, padx=5)
        ttk.Button(left_btn_frame, text="仅生成打印材料", command=self.run_print_only).pack(side=tk.LEFT, padx=5)
        ttk.Button(left_btn_frame, text="生成示例文件", command=self.create_example_files).pack(side=tk.LEFT, padx=5)

        # 右侧按钮组
        right_btn_frame = ttk.Frame(action_frame)
        right_btn_frame.pack(side=tk.RIGHT)
        ttk.Button(right_btn_frame, text="清空日志", command=self.clear_log).pack(side=tk.RIGHT, padx=5)

        # 进度条
        self.progress_var = tk.DoubleVar()
        self.progress_bar = ttk.Progressbar(log_frame, variable=self.progress_var, maximum=100)
        self.progress_bar.pack(fill=tk.X, pady=5)

        # 日志文本框
        self.log_text = scrolledtext.ScrolledText(log_frame, wrap=tk.WORD, height=20, font=("Consolas", 9))
        self.log_text.pack(fill=tk.BOTH, expand=True, pady=5)

        # 样式
        style = ttk.Style()
        style.configure("Accent.TButton", font=("微软雅黑", 10, "bold"), foreground="#0066CC")

    # ---------- 通用UI方法 ----------
    def browse_file(self, var, filetypes):
        path = filedialog.askopenfilename(initialdir=INPUT_DIR, filetypes=filetypes)
        if path:
            var.set(path)
            self.log(f"已选择文件：{os.path.basename(path)}")

    def browse_dir(self, var):
        path = filedialog.askdirectory(initialdir=OUTPUT_DIR)
        if path:
            var.set(path)
            self.log(f"已选择输出目录：{path}")

    def log(self, message):
        time_str = datetime.datetime.now().strftime("%H:%M:%S")
        self.log_text.insert(tk.END, f"[{time_str}] {message}\n")
        self.log_text.see(tk.END)
        self.root.update_idletasks()

    def clear_log(self):
        """清空日志文本框"""
        self.log_text.delete(1.0, tk.END)
        self.log("日志已清空")

    def update_progress(self, value, max_val=100):
        if max_val == 0:
            percent = 0
        else:
            percent = int(value / max_val * 100)
        self.progress_var.set(percent)
        self.root.update_idletasks()

    # ---------- 预览刷新功能 ----------
    def refresh_preview(self):
        """读取成绩文件，计算学生总数和最后一个考场容量"""
        score_path = self.score_file.get()
        if not os.path.exists(score_path):
            self.total_students_var.set("文件不存在")
            self.last_court_capacity_var.set("无法计算")
            self.log("预览失败：成绩文件不存在")
            return
        try:
            df = pd.read_excel(score_path)
            # 清洗空姓名班级行
            df = df.dropna(subset=['姓名', '班级'])
            total = len(df)
            self.total_students_var.set(str(total))

            # 获取当前设置的固定容量总和
            caps = self.get_capacities()
            if caps is None:
                self.last_court_capacity_var.set("容量无效")
                return
            fixed_sum = sum(caps)
            last_cap = total - fixed_sum
            if last_cap < 0:
                self.last_court_capacity_var.set(f"超员{fixed_sum - total}人")
                self.log(f"预览警告：固定容量总和({fixed_sum})超过学生总数({total})，最后一个考场容量为负数")
            else:
                self.last_court_capacity_var.set(str(last_cap) + " 人")
            self.log(f"预览刷新：学生总数={total}，固定容量总和={fixed_sum}，最后一个考场={last_cap}人")
        except Exception as e:
            self.total_students_var.set("读取失败")
            self.last_court_capacity_var.set("读取失败")
            self.log(f"预览刷新失败：{str(e)}")

    # ---------- 考场容量动态行操作 ----------
    def add_capacity_row(self, court_num=None):
        if court_num is None:
            court_num = len(self.capacity_rows) + 1
        row_frame = ttk.Frame(self.capacity_container)
        row_frame.pack(fill=tk.X, pady=2)
        lbl = ttk.Label(row_frame, text=f"考场 {court_num}:", width=8)
        lbl.pack(side=tk.LEFT)
        entry = ttk.Entry(row_frame, width=10)
        entry.insert(0, "2")
        entry.pack(side=tk.LEFT, padx=5)
        delete_btn = ttk.Button(row_frame, text="删除", command=lambda: self.delete_capacity_row(row_frame))
        delete_btn.pack(side=tk.LEFT, padx=5)
        self.capacity_rows.append((row_frame, entry, delete_btn, court_num))

    def delete_capacity_row(self, row_frame):
        for i, (f, e, btn, num) in enumerate(self.capacity_rows):
            if f == row_frame:
                f.destroy()
                self.capacity_rows.pop(i)
                self.renumber_courts()
                break

    def renumber_courts(self):
        for idx, (frame, entry, btn, _) in enumerate(self.capacity_rows):
            new_num = idx + 1
            for child in frame.winfo_children():
                if isinstance(child, ttk.Label) and child.cget('text').startswith('考场'):
                    child.config(text=f"考场 {new_num}:")
                    break
            self.capacity_rows[idx] = (frame, entry, btn, new_num)

    def get_capacities(self):
        caps = []
        for frame, entry, btn, num in self.capacity_rows:
            try:
                val = int(entry.get())
                if val <= 0:
                    raise ValueError
                caps.append(val)
            except:
                messagebox.showerror("错误", f"考场 {num} 的容量必须为正整数")
                return None
        return caps

    def import_capacity_from_template(self):
        template_path = self.capacity_template.get().strip()
        if not template_path or not os.path.exists(template_path):
            messagebox.showerror("错误", "请先选择有效的容量模板文件")
            return
        try:
            df = pd.read_excel(template_path, sheet_name="容量")
            df = df.dropna(subset=['容量'])
            df['容量'] = pd.to_numeric(df['容量'], errors='coerce')
            df = df.dropna(subset=['容量'])
            caps = df['容量'].astype(int).tolist()
            if not caps:
                messagebox.showerror("错误", "模板中未读取到有效的考场容量数据")
                return
            for frame, _, _, _ in self.capacity_rows:
                frame.destroy()
            self.capacity_rows.clear()
            for i, cap in enumerate(caps):
                self.add_capacity_row(i + 1)
                self.capacity_rows[-1][1].delete(0, tk.END)
                self.capacity_rows[-1][1].insert(0, str(cap))
            messagebox.showinfo("成功", f"已导入 {len(caps)} 个考场容量")
        except Exception as e:
            messagebox.showerror("错误", f"导入失败：{str(e)}")

    def export_capacity_template(self):
        export_path = self.capacity_template.get().strip()
        if not export_path:
            messagebox.showerror("错误", "请先填写导出路径")
            return
        caps = self.get_capacities()
        if caps is None:
            return
        target_dir = os.path.dirname(export_path)
        if not os.path.exists(target_dir):
            os.makedirs(target_dir)
        if os.path.exists(export_path):
            if not messagebox.askyesno("确认覆盖", f"文件已存在：\n{export_path}\n是否覆盖？"):
                return
        df = pd.DataFrame({"考场": list(range(1, len(caps) + 1)), "容量": caps})
        try:
            with pd.ExcelWriter(export_path, engine='openpyxl') as writer:
                df.to_excel(writer, index=False, sheet_name="容量")
                worksheet = writer.sheets["容量"]
                worksheet.cell(row=1, column=3, value="⚠️ 核心规则说明")
                worksheet.cell(row=2, column=3, value="1. 请在此填写前N个考场的固定容量")
                worksheet.cell(row=3, column=3, value="2. 最后一个考场自动计算，无需填写")
                worksheet.cell(row=4, column=3, value="3. 请勿修改列名，仅修改容量数值")
            messagebox.showinfo("成功", f"已导出至：{export_path}")
        except Exception as e:
            messagebox.showerror("错误", f"导出失败：{str(e)}")

    # ---------- 核心执行方法 ----------
    def run_arrange_only(self):
        """仅执行排考，生成学生考试信息.xlsx"""
        self.log("=" * 60)
        self.log("开始执行排考...")
        caps = self.get_capacities()
        if caps is None or len(caps) == 0:
            messagebox.showerror("错误", "请至少设置一个有效考场容量")
            return
        success, student_list, error = ExamArrangeCore.generate_student_exam_info(
            score_file=self.score_file.get(),
            exam_prefix=self.exam_prefix.get(),
            capacities=caps,
            output_student_file=self.student_info_output,
            output_capacity_file=self.capacity_output,
            log_callback=self.log
        )
        if success:
            self.log("✅ 排考完成！")
            messagebox.showinfo("成功",
                                f"学生考试信息已生成：\n{self.student_info_output}\n考场容量汇总：\n{self.capacity_output}")
            # 刷新预览
            self.refresh_preview()
        else:
            self.log(f"❌ 排考失败：{error}")
            messagebox.showerror("排考失败", error)

    def run_print_only(self):
        """仅生成打印材料（要求学生考试信息已存在）"""
        self.log("=" * 60)
        self.log("开始生成打印材料...")
        student_file = self.student_info_output
        if not os.path.exists(student_file):
            messagebox.showerror("错误", f"学生考试信息文件不存在：\n{student_file}\n请先执行排考或检查文件路径。")
            return
        door_file = self.door_file.get()
        if not os.path.exists(door_file):
            messagebox.showerror("错误",
                                 f"门贴数据源文件不存在：\n{door_file}\n请检查文件路径或点击【生成示例文件】创建示例。")
            return
        output_dir = self.output_dir.get()
        ensure_dir(output_dir)
        title = self.exam_title.get()
        try:
            # 读取学生列表
            wb = openpyxl.load_workbook(student_file, data_only=True)
            ws = wb.active
            student_list = []
            for row in ws.iter_rows(min_row=2, values_only=True):
                if row[0] and row[1]:
                    student_list.append({
                        '准考证号': str(row[0]).strip(),
                        '姓名': str(row[1]).strip(),
                        '班级': str(row[2]).strip().zfill(2),
                        '考场号': str(row[3]).strip().zfill(2),
                        '座号': str(row[4]).strip().zfill(2)
                    })
            doors = PrintMaterialCore.read_door_data(door_file)
            student_range = PrintMaterialCore.read_student_range(student_file)
            self.log("正在生成门贴...")
            PrintMaterialCore.generate_door_placards(doors, student_range, output_dir, title, self.log)
            self.log("正在生成座位标签...")
            PrintMaterialCore.generate_seat_labels(student_list, output_dir, self.log)
            self.log("正在生成班级考场信息...")
            PrintMaterialCore.generate_class_info(student_list, output_dir, title, self.log)
            self.log("✅ 打印材料生成完成！")
            messagebox.showinfo("完成", f"所有打印材料已生成至：\n{output_dir}")
            os.startfile(output_dir)
        except Exception as e:
            self.log(f"❌ 生成失败：{str(e)}")
            messagebox.showerror("错误", f"生成打印材料失败：{str(e)}")

    def run_full(self):
        """一键生成：先排考，再生成打印材料"""
        self.log("=" * 60)
        self.log("开始一键生成全部材料...")
        caps = self.get_capacities()
        if caps is None or len(caps) == 0:
            messagebox.showerror("错误", "请至少设置一个有效考场容量")
            return
        success, student_list, error = ExamArrangeCore.generate_student_exam_info(
            score_file=self.score_file.get(),
            exam_prefix=self.exam_prefix.get(),
            capacities=caps,
            output_student_file=self.student_info_output,
            output_capacity_file=self.capacity_output,
            log_callback=self.log
        )
        if not success:
            self.log(f"❌ 排考失败，终止后续生成。错误：{error}")
            messagebox.showerror("排考失败", error)
            return
        self.log("排考完成，开始生成打印材料...")
        door_file = self.door_file.get()
        if not os.path.exists(door_file):
            messagebox.showerror("错误",
                                 f"门贴数据源文件不存在：\n{door_file}\n请检查文件路径或点击【生成示例文件】创建示例。")
            return
        output_dir = self.output_dir.get()
        ensure_dir(output_dir)
        title = self.exam_title.get()
        try:
            # 重新读取学生列表
            wb = openpyxl.load_workbook(self.student_info_output, data_only=True)
            ws = wb.active
            student_list = []
            for row in ws.iter_rows(min_row=2, values_only=True):
                if row[0] and row[1]:
                    student_list.append({
                        '准考证号': str(row[0]).strip(),
                        '姓名': str(row[1]).strip(),
                        '班级': str(row[2]).strip().zfill(2),
                        '考场号': str(row[3]).strip().zfill(2),
                        '座号': str(row[4]).strip().zfill(2)
                    })
            doors = PrintMaterialCore.read_door_data(door_file)
            student_range = PrintMaterialCore.read_student_range(self.student_info_output)
            self.log("正在生成门贴...")
            PrintMaterialCore.generate_door_placards(doors, student_range, output_dir, title, self.log)
            self.log("正在生成座位标签...")
            PrintMaterialCore.generate_seat_labels(student_list, output_dir, self.log)
            self.log("正在生成班级考场信息...")
            PrintMaterialCore.generate_class_info(student_list, output_dir, title, self.log)
            self.log("=" * 60)
            self.log("✅ 全部材料生成完成！")
            messagebox.showinfo("完成",
                                f"所有材料已生成！\n学生考试信息：{self.student_info_output}\n打印材料输出目录：{output_dir}")
            os.startfile(output_dir)
        except Exception as e:
            self.log(f"❌ 生成打印材料失败：{str(e)}")
            messagebox.showerror("错误", f"生成打印材料失败：{str(e)}")

    def create_example_files(self):
        """生成所有示例文件"""
        try:
            create_example_score_template()
            create_example_capacity_template()
            create_example_door_file()
            self.log("✅ 示例文件已生成：")
            self.log(f"   - 成绩模板：{self.default_score_file}")
            self.log(f"   - 容量模板：{self.default_capacity_template}")
            self.log(f"   - 门贴数据源：{self.default_door_file}")
            messagebox.showinfo("成功", "示例文件已生成，请根据需要修改后再使用。")
        except Exception as e:
            self.log(f"❌ 生成示例文件失败：{str(e)}")
            messagebox.showerror("错误", f"生成示例文件失败：{str(e)}")


# ---------------------------- 程序入口 ----------------------------
def main():
    root = tk.Tk()
    app = UnifiedExamTool(root)
    root.mainloop()


if __name__ == "__main__":
    main()
